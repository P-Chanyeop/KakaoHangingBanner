# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

doc = Document()

# 제목
title = doc.add_heading('', 0)
run = title.add_run('웹 애플리케이션 개발 견적서')
run.font.size = Pt(24)
run.font.name = '맑은 고딕'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 부제목
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('참신한 게시대 - 현수막 게시대 관리 시스템')
run.font.size = Pt(14)
run.font.name = '맑은 고딕'

doc.add_paragraph()

# 기본 정보
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('프로젝트명', '참신한 게시대 - 현수막 게시대 관리 시스템'),
    ('작성일', datetime.now().strftime('%Y년 %m월 %d일')),
    ('견적 유효기간', '발행일로부터 30일'),
    ('개발 기간', '약 4~6주 (협의 가능)')
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    set_cell_shading(info_table.rows[i].cells[0], 'E8E8E8')

doc.add_paragraph()
doc.add_heading('1. 프로젝트 개요', level=1)
doc.add_paragraph('경상북도와 경상남도 지역의 현수막 게시대 위치를 관리하고, 지도에서 쉽게 찾을 수 있는 웹 애플리케이션입니다.')

doc.add_heading('2. 기술 스택', level=1)
tech_para = doc.add_paragraph()
tech_para.add_run('Backend: ').bold = True
tech_para.add_run('Spring Boot 3.5.3, Java 17, Spring Security, JWT, JPA, MySQL\n')
tech_para.add_run('Frontend: ').bold = True
tech_para.add_run('React 18.3.1, React Router 6, Leaflet, Kakao Maps, Naver Maps\n')
tech_para.add_run('외부 API: ').bold = True
tech_para.add_run('카카오맵 API, 네이버맵 API, VWorld API')

doc.add_heading('3. 상세 기능 및 견적', level=1)

# 견적 테이블
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 헤더
header_cells = table.rows[0].cells
headers = ['구분', '기능 상세', '작업일', '금액(원)']
for i, header in enumerate(headers):
    header_cells[i].text = header
    set_cell_shading(header_cells[i], '4472C4')
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = None
            run.font.bold = True

# 데이터
items = [
    ('1. 핵심 시스템', '사용자 인증 (JWT)\n다중 지도 시스템\n게시대 관리 (CRUD)\n마커/라벨 시스템\n역지오코딩', '', '1,600,000'),
    ('2. 로드뷰 기능', '네이버 로드뷰 연동\n선택/토글 모드\n미니맵 연동', '', '100,000'),
    ('3. 관리자 페이지', '버튼 링크 관리\nHero 이미지 관리\n팝업 메시지 관리', '', '150,000'),
    ('4. 캘린더 기능', '월별 일정 관리\n일정 CRUD\n색상 코드 구분', '', '50,000'),
    ('5. 장소 검색 기능', '카카오 Places API 연동\n키워드 검색\n검색 결과 드롭다운', '', '100,000'),
]

for item in items:
    row = table.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

# 합계
total_row = table.add_row().cells
total_row[0].text = '합계'
total_row[1].text = ''
total_row[2].text = ''
total_row[3].text = '2,000,000'
for cell in total_row:
    set_cell_shading(cell, 'FFF2CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

doc.add_paragraph()
doc.add_heading('4. 결제 조건', level=1)
payment = doc.add_paragraph()
payment.add_run('• 계약금: ').bold = True
payment.add_run('30% (계약 시)\n')
payment.add_run('• 중도금: ').bold = True
payment.add_run('40% (개발 50% 완료 시)\n')
payment.add_run('• 잔금: ').bold = True
payment.add_run('30% (최종 납품 시)')

doc.add_heading('5. 포함 사항', level=1)
doc.add_paragraph('• 소스코드 전체 제공\n• 설치 및 배포 가이드 문서\n• 1개월 무상 유지보수\n• 사용자 매뉴얼')

doc.add_heading('6. 불포함 사항', level=1)
doc.add_paragraph('• 서버 호스팅 비용\n• 도메인 비용\n• 외부 API 사용료 (카카오맵, 네이버맵 등)\n• 1개월 이후 유지보수 (별도 협의)')

doc.add_heading('7. 특이사항', level=1)
doc.add_paragraph('• 본 견적은 현재 구현된 기능 기준입니다.\n• 추가 기능 요청 시 별도 견적이 필요합니다.\n• 개발 기간은 요구사항 확정 후 조정될 수 있습니다.')

doc.add_paragraph()
doc.add_paragraph()

# 서명란
sign_table = doc.add_table(rows=3, cols=2)
sign_table.style = 'Table Grid'
sign_data = [
    ('발주처', '수주처'),
    ('상호:', '상호:'),
    ('대표:', '대표:'),
]
for i, (left, right) in enumerate(sign_data):
    sign_table.rows[i].cells[0].text = left
    sign_table.rows[i].cells[1].text = right
    if i == 0:
        set_cell_shading(sign_table.rows[i].cells[0], 'E8E8E8')
        set_cell_shading(sign_table.rows[i].cells[1], 'E8E8E8')

# 저장
doc.save('/mnt/c/GITHUB/카카오행잉배너/KakaoHangingBanner/개발견적서_참신한게시대.docx')
print('견적서 생성 완료: 개발견적서_참신한게시대.docx')
